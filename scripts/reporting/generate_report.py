from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from PIL import Image, ImageDraw, ImageFont

# --- Content Definitions (Expanded) ---

PROJECT_TITLE = "QNA-AUTH: QUANTUM NOISE ASSISTED AUTHENTICATION"

ABSTRACT_TEXT = (
    "The proliferation of Internet of Things (IoT) devices and mobile computing has exposed significant "
    "vulnerabilities in traditional authentication mechanisms. Static credentials, such as passwords and "
    "cryptographic keys, are susceptible to theft, phishing, and replay attacks. Furthermore, standard "
    "biometrics (fingerprint, facial recognition) rely on immutable physical traits that, once compromised, "
    "cannot be reset. To address these critical security gaps, this project proposes QNA-Auth "
    "(Quantum Noise Assisted Authentication), a novel framework that derives non-invertible device "
    "fingerprints from high-entropy physical noise sources.\n\n"
    "QNA-Auth shifts the paradigm from 'what you know' to 'what your hardware is' by leveraging transient, "
    "stochastic physical properties. The system aggregates entropy from three distinct sources: quantum "
    "vacuum fluctuations (via QRNG APIs), thermal sensor noise from camera sensors (Dark Frames), and "
    "electromagnetic/ambient noise from microphones. A robust preprocessing pipeline extracts statistical, "
    "spectral (FFT), and complexity-based features from this raw noise. These features are then mapped into "
    "a high-dimensional embedding space using a Siamese Neural Network trained via contrastive learning. "
    "This deep learning approach ensures that noise samples from the same device cluster tightly together, "
    "while samples from different devices are separated by a distinct margin.\n\n"
    "The system is implemented as a full-stack prototype utilizing Python (FastAPI) for the backend, "
    "PyTorch for the machine learning engine, and React for the client interface. Crucially, the system "
    "implements a Challenge-Response Protocol using cryptographic nonces to prevent replay attacks, ensuring "
    "that every authentication attempt requires fresh, live noise data. This project demonstrates that "
    "commodity hardware can be uniquely identified without specialized chips, offering a scalable, "
    "privacy-preserving primitive for next-generation device security."
)

CHAPTERS = [
    {
        "title": "INTRODUCTION",
        "sections": [
            {
                "heading": "BACKGROUND",
                "content": (
                    "In the modern digital ecosystem, the integrity of a system is defined by its ability to "
                    "distinguish trusted entities from malicious actors. Historically, this has been achieved "
                    "through three factors of authentication: knowledge (passwords), possession (tokens), and "
                    "inherence (biometrics). However, the cybersecurity landscape has shifted dramatically. "
                    "Knowledge-based secrets are routinely compromised through massive database breaches. "
                    "Possession-based tokens can be cloned or stolen. Even biometric data, often considered "
                    "the gold standard, faces threats from DeepFake technology and high-resolution sensor spoofing.\n\n"
                    "QNA-Auth introduces a physics-based approach to this problem. Every electronic device, "
                    "due to manufacturing variations in silicon, exhibits unique noise characteristics—often "
                    "referred to as a Physical Unclonable Function (PUF). By combining these intrinsic "
                    "hardware imperfections with Quantum Random Number Generation (QRNG), we can generate "
                    "high-quality entropy that acts as a unique fingerprint. Unlike a stored key, this "
                    "fingerprint is generated dynamically only when needed and does not persist in memory, "
                    "making it inherently resistant to static data theft."
                )
            },
            {
                "heading": "MOTIVATION",
                "content": (
                    "The primary motivation for this project stems from the inherent weakness of 'reproducible "
                    "credentials.' If an attacker gains access to a server's database and steals a password "
                    "hash or a biometric template, they can theoretically masquerade as the user indefinitely. "
                    "This 'static' nature of current authentication is a single point of failure.\n\n"
                    "There is a critical need for an authentication mechanism where the 'key' is:\n"
                    "1. Transient: It exists only during the moment of authentication.\n"
                    "2. Non-Invertible: Even if the stored data is stolen, it cannot be used to reverse-engineer "
                    "the original key.\n"
                    "3. Hardware-Bound: It ties the identity to a specific physical device, preventing remote "
                    "emulation attacks.\n\n"
                    "Existing solutions like hardware security modules (HSMs) or TPMs are expensive and "
                    "proprietary. This project aims to democratize this level of security by using machine "
                    "learning to extract 'soft PUFs' from standard, commodity hardware sensors found in "
                    "everyday laptops and smartphones."
                )
            },
            {
                "heading": "SCOPE OF THE PROJECT",
                "content": (
                    "The scope of QNA-Auth extends to the design, development, and validation of a complete "
                    "authentication framework.\n\n"
                    "In Scope:\n"
                    "- Multi-Source Data Collection: Implementation of modules to capture noise from Webcams "
                    "(thermal noise), Microphones (ambient/EM noise), and external Quantum RNG APIs.\n"
                    "- Advanced Signal Processing: Development of a pipeline to clean, normalize, and extract "
                    "features (Fourier Transforms, Entropy) from raw noisy signals.\n"
                    "- Deep Metric Learning: Training a Siamese Neural Network to learn a distance function "
                    "that quantifies the similarity between noise samples.\n"
                    "- Full-Stack Implementation: A functional prototype with a RESTful API backend and a "
                    "responsive web-based frontend.\n"
                    "- Security Protocol Design: Implementation of a nonce-based challenge-response flow to "
                    "mitigate replay attacks.\n\n"
                    "Out of Scope:\n"
                    "- User Behavioral Biometrics: The project focuses strictly on device fingerprinting, "
                    "not user behavior (e.g., typing patterns).\n"
                    "- Hardware Fabrication: We utilize existing commodity sensors; we are not designing new "
                    "hardware chips.\n"
                    "- Large-scale Cloud Deployment: The prototype is designed for local and network testing "
                    "environments."
                )
            }
        ]
    },
    {
        "title": "PROJECT DESCRIPTION AND GOALS",
        "sections": [
            {
                "heading": "LITERATURE REVIEW",
                "subsections": [
                    {
                        "heading": "Machine Learning Based Authentication",
                        "content": (
                            "The intersection of Machine Learning (ML) and cybersecurity has been a fertile "
                            "ground for research. Traditional approaches often utilize Support Vector Machines "
                            "(SVMs) or Random Forests to classify network traffic anomalies or user behavioral "
                            "patterns. For instance, researchers have successfully used ML to identify devices "
                            "based on clock skew and TCP timestamp variations. However, these methods often "
                            "rely on 'hand-crafted' features which may not capture the subtle, non-linear "
                            "complexities of sensor noise data. Furthermore, traditional classifiers require "
                            "retraining whenever a new device is added to the system, limiting their scalability."
                        )
                    },
                    {
                        "heading": "Deep Learning Based Approaches",
                        "content": (
                            "Deep Learning has revolutionized pattern recognition. In the context of "
                            "authentication, Metric Learning approaches, particularly Siamese Networks, have "
                            "gained prominence. Originally popularized by systems like FaceNet for facial "
                            "recognition, these networks do not classify inputs into fixed categories. "
                            "Instead, they map inputs to a continuous vector space where distance equals "
                            "similarity. QNA-Auth adapts this architecture for the domain of signal processing. "
                            "By treating sensor noise as a high-dimensional signal, we employ Deep Neural "
                            "Networks (DNNs) to learn a robust embedding space. This allows for 'few-shot "
                            "learning,' where a device can be enrolled with only a small number of samples, "
                            "and the system can recognize it without retraining the entire model."
                        )
                    }
                ]
            },
            {
                "heading": "GAPS IDENTIFIED",
                "content": (
                    "Despite advancements, significant gaps remain in current literature and industry practice:\n"
                    "1. Reliance on Volatile Secrets: Most 'secure' enclave solutions still rely on storing a "
                    "private key. If the hardware is decapped or subjected to side-channel power analysis, "
                    "these keys can be extracted.\n"
                    "2. Lack of Open Implementations: Most PUF-based research is theoretical or tied to "
                    "proprietary FPGA implementations, making it inaccessible for general software developers.\n"
                    "3. Vulnerability to Replay: Simple noise-based identification systems often fail to "
                    "account for an attacker simply recording the noise signal and playing it back.\n"
                    "4. Static Biometrics: Current biometric systems lack a 'cancelable' feature. If your "
                    "fingerprint is stolen, you cannot change it."
                )
            },
            {
                "heading": "OBJECTIVES",
                "content": (
                    "The specific objectives of QNA-Auth are as follows:\n"
                    "1. To develop a robust Noise Collection Pipeline capable of interfacing with standard "
                    "OS hardware (Camera/Mic) to capture high-entropy raw data.\n"
                    "2. To implement a Feature Extraction Engine that computes spectral (FFT), statistical, "
                    "and complexity-based features to reduce noise dimensionality.\n"
                    "3. To design and train a Siamese Neural Network using Triplet Loss to creating unique, "
                    "distinguishable embeddings for different devices.\n"
                    "4. To implement a Cryptographically Secure Challenge-Response Protocol that utilizes "
                    "server-generated nonces to prevent replay attacks.\n"
                    "5. To build a User-Friendly Interface (React) and a scalable API (FastAPI) to "
                    "demonstrate the practicality of the system."
                )
            },
            {
                "heading": "PROBLEM STATEMENT",
                "content": (
                    "In the domain of device authentication, reliance on static secrets (keys, passwords) "
                    "creates a permanent vulnerability window. If a secret is stolen, the device identity is "
                    "compromised. There is currently no widely available, software-defined solution that can "
                    "uniquely authenticate commodity devices based on their intrinsic physical properties "
                    "without requiring specialized hardware add-ons. The challenge is to convert the "
                    "stochastic, noisy output of standard sensors into a stable, reliable digital fingerprint "
                    "that is robust against replication and replay attacks."
                )
            }
        ]
    },
    {
        "title": "TECHNICAL SPECIFICATION",
        "sections": [
            {
                "heading": "REQUIREMENTS",
                "subsections": [
                    {
                        "heading": "Functional Requirements",
                        "content": (
                            "- Enrollment: The system must allow a user to register a new device by capturing "
                            "20-50 noise samples.\n"
                            "- Authentication: The system must verify a registered device by capturing fresh "
                            "noise samples and comparing them against the stored profile.\n"
                            "- Hardware Interface: The application must successfully access the host machine's "
                            "webcam and microphone driver layers.\n"
                            "- Quantum Integration: The system must fetch real-time quantum entropy from the "
                            "ANU QRNG API to salt the local sensor noise.\n"
                            "- Security Checks: The system must validate cryptographic nonces and reject any "
                            "authentication attempt that uses an expired or previously used nonce."
                        )
                    },
                    {
                        "heading": "Non-Functional Requirements",
                        "content": (
                            "- Latency: The end-to-end authentication process (capture -> process -> inference "
                            "-> verify) should complete within 3 seconds to ensure good UX.\n"
                            "- Accuracy: The model should achieve a False Acceptance Rate (FAR) of < 0.1% to "
                            "prevent unauthorized access.\n"
                            "- Reliability: The system should handle sensor initialization failures gracefully "
                            "with fallback mechanisms.\n"
                            "- Privacy: Raw sensor data (photos/audio) must never be stored persistently; "
                            "only the mathematical embeddings should be saved."
                        )
                    }
                ]
            },
            {
                "heading": "FEASIBILITY STUDY",
                "content": (
                    "Technical Feasibility: The project is technically feasible as it relies on mature, "
                    "open-source libraries. Python provides robust support for hardware interaction (OpenCV, "
                    "PyAudio/SoundDevice) and machine learning (PyTorch). The 'Siamese Network' architecture "
                    "is well-documented and proven for similarity tasks.\n\n"
                    "Economic Feasibility: The solution is highly cost-effective. Unlike hardware tokens "
                    "(YubiKeys) or Smart Cards which cost money per unit, QNA-Auth utilizes hardware that "
                    "users already possess. The software stack consists entirely of free, open-source "
                    "components, incurring zero licensing costs.\n\n"
                    "Social Feasibility: The system is socially viable as it enhances privacy. Users are often "
                    "wary of biometric systems storing photos of their faces. QNA-Auth captures 'Dark Frames' "
                    "(lens covered) and ambient silence, meaning no identifiable personal information (PII) "
                    "like facial features or conversations are ever recorded or stored."
                )
            },
            {
                "heading": "SYSTEM SPECIFICATION",
                "subsections": [
                    {
                        "heading": "Hardware Specification",
                        "content": (
                            "- Server: Intel Core i5/i7 or AMD Ryzen 5 (for model inference).\n"
                            "- RAM: 8GB minimum (16GB recommended for training).\n"
                            "- Client Device: Any device with a functional Webcam (min 720p resolution) and Microphone.\n"
                            "- GPU (Optional): NVIDIA GTX/RTX series for accelerated model training (CUDA)."
                        )
                    },
                    {
                        "heading": "Software Specification",
                        "content": (
                            "- Operating System: Cross-platform (Windows, Linux, macOS).\n"
                            "- Backend Framework: FastAPI (Python 3.9+).\n"
                            "- Frontend Framework: React 18 (TypeScript) with Vite build tool.\n"
                            "- ML Engine: PyTorch 2.0+.\n"
                            "- Computer Vision: OpenCV (cv2).\n"
                            "- Signal Processing: NumPy, SciPy."
                        )
                    }
                ]
            }
        ]
    },
    {
        "title": "DESIGN APPROACH AND DETAILS",
        "sections": [
            {
                "heading": "SYSTEM ARCHITECTURE",
                "content": (
                    "The QNA-Auth system employs a Microservices-oriented Client-Server Architecture.\n\n"
                    "1. The Client (Frontend): Developed in React, this layer acts as the data acquisition unit. "
                    "It interacts with the browser's MediaStream API to capture raw video frames and audio buffers.\n"
                    "2. The API Gateway (FastAPI): This serves as the central orchestrator. It exposes REST "
                    "endpoints (/enroll, /authenticate, /challenge) and manages the flow of data.\n"
                    "3. The Intelligence Layer (ML Model): A PyTorch-based service that accepts raw noise vectors "
                    "and outputs 128-dimensional L2-normalized embeddings.\n"
                    "4. The Storage Layer: A file-system-based storage (simulating a database) that holds the "
                    "serialized device embeddings (.pt files) and device metadata (.json)."
                ),
                "image": "architecture_diagram.png" 
            },
            {
                "heading": "DESIGN",
                "subsections": [
                    {
                        "heading": "Data Flow Diagram",
                        "content": (
                            "The data flow for a typical authentication request is as follows:\n"
                            "1. Request Challenge: Client requests a login challenge. Server generates a random nonce.\n"
                            "2. Data Capture: Client captures physical noise (Camera/Mic) + Quantum Noise (API).\n"
                            "3. Transmission: Client sends the Noise Data + Signed Nonce to the Server.\n"
                            "4. Preprocessing: Server validates Nonce. Noise data goes through FFT and Statistical analysis.\n"
                            "5. Inference: Extracted features are passed to the Siamese Network. The network outputs a Live_Embedding.\n"
                            "6. Verification: Server loads the Stored_Embedding for the claimed Device ID.\n"
                            "7. Comparison: Cosine Similarity is calculated.\n"
                            "8. Decision: If Similarity > Threshold (0.85), Access Granted."
                        )
                    },
                    {
                        "heading": "Class Diagram",
                        "content": (
                            "Key Object-Oriented classes governing the system:\n"
                            "- DeviceEnroller: Responsible for the multi-step process of collecting samples, "
                            "aggregating them into a centroid, and persisting the profile to disk.\n"
                            "- DeviceAuthenticator: Handles the logic of comparing a fresh incoming sample "
                            "against a stored profile.\n"
                            "- NoiseCollector (Interface): An abstract base class implemented by "
                            "CameraNoiseCollector and MicrophoneNoiseCollector to ensure uniform data acquisition.\n"
                            "- SiameseNetwork (nn.Module): The PyTorch neural network class defining the layers.\n"
                            "- ChallengeResponseProtocol: Manages the lifecycle of cryptographic nonces, "
                            "including generation, storage, and expiry validation."
                        )
                    }
                ]
            }
        ]
    },
    {
        "title": "METHODOLOGY AND TESTING",
        "sections": [
            {
                "heading": "MODULE DESCRIPTION",
                "content": (
                    "1. Noise Collection Module: Interfaces directly with hardware drivers. "
                    "Captures 'Dark Frames' by reading sensor data with the lens covered. "
                    "Records ambient silence to analyze the microphone's self-noise floor. "
                    "Fetches live quantum data from the ANU QRNG API.\n\n"
                    "2. Preprocessing Module: Transforms raw arrays into compact feature vectors. "
                    "Calculates Statistical Features (Mean, Variance, Skewness, Kurtosis), "
                    "Spectral Features (FFT bands), and Complexity Features (Shannon Entropy).\n\n"
                    "3. Model Module (Siamese Network): The core intelligence consisting of two identical "
                    "subnetworks. It uses Triplet Loss to minimize the distance between same-device samples "
                    "while maximizing the distance between different devices.\n\n"
                    "4. Auth Module (Challenge-Response): To prevent replay attacks, the server generates a "
                    "random nonce. The client must return the noise data bundled with a hash of the nonce, "
                    "ensuring timeliness."
                )
            },
            {
                "heading": "TESTING",
                "content": (
                    "Testing was conducted in three phases:\n\n"
                    "Unit Testing: Individual modules were tested in isolation using pytest. "
                    "Verified hardware access and correct feature extraction shapes.\n\n"
                    "Integration Testing: The full API flow was tested using test_enrollment.py, "
                    "simulating client enrollment and authentication sequences.\n\n"
                    "Cross-Device Validation: Performed experiments enrolling Device A and attempting to "
                    "authenticate using noise from Device B. "
                    "Intra-device similarity (A vs A) consistently scored > 0.90. "
                    "Inter-device similarity (A vs B) consistently scored < 0.30, confirming distinguishability."
                )
            }
        ]
    }
]

REFERENCES_TEXT = (
    "Apruzzese, G., Laskov, P., Montes de Oca, E., Mallouli, W., Brdalo Rapa, L., Grammatopoulos, A. V., & Di Franco, F. (2023). The role of machine learning in cybersecurity. Digital Threats: Research and Practice, 4(1), 1-38.\n"
    "Kumar, S., Gupta, U., Singh, A. K., & Singh, A. K. (2023). AI: Revolutionizing cyber security in the Digital Era. J. Comput. Mech. Manag, 2(3), 31-42.\n"
    "Nisbet, R., Miner, G. D., & McCormick, K. (2024). Handbook of Statistical Analysis: AI and ML Applications. Elsevier.\n"
    "Schroff, F., Kalenichenko, D., & Philbin, J. (2015). FaceNet: A unified embedding for face recognition and clustering. CVPR.\n"
    "Documentation. (2024). PyTorch Documentation. Retrieved from https://pytorch.org/docs/\n"
    "Documentation. (2024). FastAPI Documentation. Retrieved from https://fastapi.tiangolo.com/\n"
    "Australian National University (ANU). (2024). Quantum Random Numbers Server API. Retrieved from https://qrng.anu.edu.au/"
)

# --- Helper Functions ---

def create_architecture_diagram():
    """Generates a simple architecture diagram image."""
    img = Image.new('RGB', (800, 400), color='white')
    draw = ImageDraw.Draw(img)
    
    # Draw simple boxes
    # Client
    draw.rectangle([50, 150, 200, 250], outline="black", width=2)
    draw.text((100, 190), "Client\n(React UI)", fill="black")
    
    # Arrow
    draw.line([200, 200, 300, 200], fill="black", width=2)
    draw.polygon([(300, 200), (290, 195), (290, 205)], fill="black")
    
    # API
    draw.rectangle([300, 150, 450, 250], outline="black", width=2)
    draw.text((350, 190), "Backend\n(FastAPI)", fill="black")
    
    # Arrow
    draw.line([450, 200, 550, 200], fill="black", width=2)
    draw.polygon([(550, 200), (540, 195), (540, 205)], fill="black")
    
    # Model
    draw.rectangle([550, 150, 700, 250], outline="black", width=2)
    draw.text((600, 190), "ML Model\n(PyTorch)", fill="black")
    
    img.save("architecture_diagram.png")
    return "architecture_diagram.png"

def setup_styles(doc):
    """Configures styles to match the template."""
    styles = doc.styles
    
    # Normal Text (12pt Times New Roman)
    style_normal = styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style_normal.paragraph_format
    paragraph_format.line_spacing = 1.5
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Heading 1 (Level 0 - Chapter Titles)
    style_h1 = styles['Heading 1']
    font = style_h1.font
    font.name = 'Times New Roman'
    font.size = Pt(16)
    font.bold = True
    font.all_caps = True
    style_h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_h1.paragraph_format.space_after = Pt(24)
    
    # Heading 2 (Level 1 - Main Headings 1.1)
    style_h2 = styles['Heading 2']
    font = style_h2.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    font.bold = True
    font.all_caps = True
    style_h2.paragraph_format.space_before = Pt(18)
    style_h2.paragraph_format.space_after = Pt(12)

    # Heading 3 (Level 2 - Sub Headings 1.1.1)
    style_h3 = styles['Heading 3']
    font = style_h3.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.bold = True
    style_h3.paragraph_format.space_before = Pt(12)
    style_h3.paragraph_format.space_after = Pt(6)

def add_cover_page(doc):
    """Creates the cover page according to VIT template."""
    
    # Course Code
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BCSE498J Project-II / CBS1904 - Capstone Project")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.font.bold = True
    
    # Project Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    p.paragraph_format.space_after = Pt(48)
    run = p.add_run(PROJECT_TITLE)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    # Students
    students = [
        "Reg. No. 22BCI00001, STUDENT NAME 1",
        "Reg. No. 22BCI00002, STUDENT NAME 2",
        "Reg. No. 22BCI00003, STUDENT NAME 3"
    ]
    for stud in students:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(stud)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
    
    # Supervision
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    run = p.add_run("Under the Supervision of")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.italic = True 
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run("Prof. GUIDE NAME\nDesignation\nSchool of Computer Science and Engineering (SCOPE)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    # Footer Section (Degree)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(48)
    text = "B.Tech.\nin\nComputer Science and Engineering\n(with specialization in Information Security)"
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("[VIT LOGO PLACEHOLDER]")
    run.font.size = Pt(12)
    
    # School and Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("School of Computer Science and Engineering (SCOPE)\nFebruary 2026")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True

    doc.add_section(WD_SECTION.NEW_PAGE)

def create_document():
    doc = Document()
    
    # Page Setup
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.left_margin = Inches(1.5)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    
    setup_styles(doc)
    
    # --- Cover Page ---
    add_cover_page(doc)
    
    # --- Abstract ---
    p = doc.add_heading('ABSTRACT', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph(ABSTRACT_TEXT)
    doc.add_page_break()
    
    # --- Table of Contents Placeholder ---
    p = doc.add_heading('TABLE OF CONTENTS', level=1)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph("[Table of Contents to be generated in Word]")
    doc.add_page_break()
    
    # --- Chapters ---
    img_path = create_architecture_diagram()
    
    for ch_idx, chapter in enumerate(CHAPTERS, 1):
        if ch_idx > 1:
            doc.add_page_break()
            
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"CHAPTER {ch_idx}")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(chapter['title'])
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        p.paragraph_format.space_after = Pt(24)
        
        for sec_idx, section in enumerate(chapter['sections'], 1):
            h_text = f"{ch_idx}.{sec_idx} {section['heading']}"
            h = doc.add_heading(h_text, level=2)
            
            if 'content' in section:
                # Handle paragraphs separated by \n\n in content
                paragraphs = section['content'].split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            
            if 'image' in section:
                try:
                    doc.add_picture(section['image'], width=Inches(5))
                    last_p = doc.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption = doc.add_paragraph(f"Figure {ch_idx}.{sec_idx}: System Architecture")
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption.style = doc.styles['Normal']
                    caption.style.font.size = Pt(10)
                except Exception:
                    pass

            if 'subsections' in section:
                for sub_idx, subsection in enumerate(section['subsections'], 1):
                    h_sub_text = f"{ch_idx}.{sec_idx}.{sub_idx} {subsection['heading']}"
                    doc.add_heading(h_sub_text, level=3)
                    
                    sub_paragraphs = subsection['content'].split('\n\n')
                    for para in sub_paragraphs:
                        if para.strip():
                            doc.add_paragraph(para.strip())
    
    # --- References ---
    doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("REFERENCES")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(16)
    run.font.bold = True
    
    doc.add_paragraph(REFERENCES_TEXT)
    
    # Save
    doc.save("QNA_Auth_Project_Report.docx")
    
    if os.path.exists(img_path):
        os.remove(img_path)

if __name__ == "__main__":
    create_document()